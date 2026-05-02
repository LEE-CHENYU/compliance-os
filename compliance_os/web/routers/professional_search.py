"""HTTP surface for professional search onboarding.

- `POST /api/professional-search` — user intake. Accepts a case brief
  (required) plus any number of uploaded files. Creates a
  `ProfessionalSearchRequestRow`, extracts text from uploads into the
  row's `uploaded_notes`, then kicks off the Anthropic-powered runner
  as a background task. Returns the request id so the client can poll.

- `GET /api/professional-search/{id}` — returns the current row
  (status, per-persona progress, tier report when complete).
"""
from __future__ import annotations

import datetime as _dt
import html as _html
import io
import json
import logging
from pathlib import Path
from typing import Any

import yaml as _yaml
from fastapi import (
    APIRouter,
    BackgroundTasks,
    Depends,
    File,
    Form,
    Header,
    HTTPException,
    Request,
    UploadFile,
)
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.orm import Session

from compliance_os.settings import settings
from compliance_os.web.models.auth import UserRow
from compliance_os.web.models.database import get_session
from compliance_os.web.models.tables import ProfessionalSearchRequestRow
from compliance_os.web.services.auth_service import get_bearer_payload

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/professional-search", tags=["professional-search"])


MAX_UPLOAD_SIZE = 10 * 1024 * 1024  # 10MB per file
MAX_TEXT_PER_FILE = 4000             # chars appended to brief per file
MAX_UPLOADS = 8


# ----------------------------- helpers --------------------------------

def _extract_text(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Best-effort text extraction. Returns '' on failure — never raises."""
    try:
        ext = Path(filename).suffix.lower()
        if ext in {".txt", ".md", ".csv"} or content_type.startswith("text/"):
            return file_bytes.decode("utf-8", errors="replace")
        if ext == ".pdf" or content_type == "application/pdf":
            try:
                import fitz  # pymupdf
            except ImportError:
                return ""
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            try:
                return "\n".join(page.get_text() for page in doc)
            finally:
                doc.close()
        if ext == ".docx":
            try:
                from docx import Document as _Docx
            except ImportError:
                return ""
            return "\n".join(p.text for p in _Docx(io.BytesIO(file_bytes)).paragraphs)
    except Exception:
        return ""
    return ""


def _build_uploaded_notes(files: list[UploadFile]) -> str:
    """Concatenate trimmed text from each upload into a single notes blob."""
    parts: list[str] = []
    for f in files[:MAX_UPLOADS]:
        raw = f.file.read(MAX_UPLOAD_SIZE + 1)
        if len(raw) > MAX_UPLOAD_SIZE:
            parts.append(f"[{f.filename}: skipped — exceeds {MAX_UPLOAD_SIZE} bytes]")
            continue
        text = _extract_text(raw, f.filename or "upload", f.content_type or "")
        text = text.strip()
        if not text:
            parts.append(f"[{f.filename}: no extractable text]")
            continue
        if len(text) > MAX_TEXT_PER_FILE:
            text = text[:MAX_TEXT_PER_FILE] + f"\n... [truncated — original {len(text)} chars]"
        parts.append(f"--- {f.filename} ---\n{text}")
    return "\n\n".join(parts)


# ----------------------------- schemas --------------------------------

class SearchResponse(BaseModel):
    id: str
    status: str
    purpose: str
    vertical: str
    case_brief: str
    uploaded_notes: str | None
    persona_status: dict
    tier_report: list | None
    error: str | None
    created_at: str
    completed_at: str | None
    paid_at: str | None
    is_paid: bool
    is_claimed: bool
    stripe_customer_email: str | None
    # Set when this search was launched from inside a case (via the
    # "Find a specialist" CTA). The frontend uses it to enable the
    # per-firm "+ Track" buttons on the results page.
    case_id: str | None
    # Stage-2 enrichment lifecycle — null fields when never dispatched.
    # Frontend polls these to decide whether to show "verifying individual
    # attorney credentials..." banner above the firm list.
    enrichment_status: str = "idle"
    enrichment_started_at: str | None = None
    enrichment_completed_at: str | None = None
    # Per-firm rich data (Stage 1 personas + Stage 2 enrichment overlays).
    # Only exposed when paid_at is set — pre-payment users see just the
    # top-5 tier_report slice, never the full contact-info list. Field is
    # always present in the response shape but null for unpaid rows.
    firms_data: list | None = None


class EnrichmentStatusResponse(BaseModel):
    """Lightweight polling shape for the /paid page — we don't want every
    poll to drag the full firms_data + tier_report payload across the wire.
    """
    status: str  # idle | enriching | complete | failed
    started_at: str | None
    completed_at: str | None
    # firm-level rough completion: how many firms have at least one
    # `_lead_attorney_*` field set. Lets the UI show a "X of Y firms
    # enriched" progress bar during the ~2-3 min window.
    firms_enriched: int
    firms_total: int


def _normalize_tier_report(rows: list | None) -> list | None:
    """Expose a stable `firm` key for attorney and non-attorney reports.

    Older CPA/bank/CAA rows were stored directly from `v_vendor_comparison`,
    whose name column is `vendor` while attorney rows use `firm`. The public
    status page expects one canonical display key, so normalize on serialize
    as well as at new-row creation time.
    """
    if rows is None:
        return None
    normalized: list = []
    for row in rows:
        if not isinstance(row, dict):
            normalized.append(row)
            continue
        item = dict(row)
        firm = item.get("firm") or item.get("vendor") or item.get("name")
        if firm is not None and not item.get("firm"):
            item["firm"] = str(firm)
        normalized.append(item)
    return normalized


def _public_persona_status(statuses: dict | None) -> dict:
    """Return only progress fields that belong in the browser payload."""
    public: dict = {}
    for persona_id, status in (statuses or {}).items():
        if not isinstance(status, dict):
            continue
        item: dict = {}
        for key in ("status", "firm_count", "started_at", "finished_at"):
            if key in status:
                item[key] = status[key]
        if status.get("status") == "skipped":
            item["reason"] = "Skipped because this search axis was not relevant to the case."
        public[persona_id] = item
    return public


def _public_firms_data(firms: list | None) -> list | None:
    if firms is None:
        return None
    public: list = []
    hidden_keys = {"_enrichment_error"}
    for firm in firms:
        if not isinstance(firm, dict):
            public.append(firm)
            continue
        public.append({
            key: value for key, value in firm.items()
            if key not in hidden_keys
        })
    return public


def _public_search_error(error: str | None) -> str | None:
    if not error:
        return None
    return (
        "This search did not complete. Please retry or contact Guardian support "
        "from your dashboard."
    )


def _serialize(
    row: ProfessionalSearchRequestRow,
    *,
    include_pii: bool = False,
) -> SearchResponse:
    """Serialize a search row to the public API shape.

    PII gating: `case_brief`, `uploaded_notes`, and `stripe_customer_email`
    contain user-supplied / Stripe-captured personal data. The status
    page polls `GET /{request_id}` *unauthenticated* by design (the URL
    is shareable + the user might not have signed up yet). Returning
    the full brief / customer email there leaks PII to anyone who
    learns the (UUID) request_id — including via Referer headers when
    the page is shared. Pass `include_pii=True` only on auth-gated
    paths (claim, mine/list, future authenticated GETs).
    """
    return SearchResponse(
        id=row.id,
        status=row.status,
        purpose=row.purpose,
        vertical=row.vertical,
        case_brief=(row.case_brief if include_pii else ""),
        uploaded_notes=(row.uploaded_notes if include_pii else None),
        persona_status=_public_persona_status(row.persona_status),
        tier_report=_normalize_tier_report(row.tier_report),
        error=_public_search_error(row.error),
        created_at=row.created_at.isoformat() if row.created_at else "",
        completed_at=row.completed_at.isoformat() if row.completed_at else None,
        paid_at=row.paid_at.isoformat() if row.paid_at else None,
        is_paid=row.paid_at is not None,
        is_claimed=row.user_id is not None,
        # Customer email is needed by the post-purchase page to pre-fill
        # the signup form, but only the user who's about to claim should
        # see it. Anonymous polling gets None.
        stripe_customer_email=(row.stripe_customer_email if include_pii else None),
        case_id=row.case_id,
        enrichment_status=row.enrichment_status or "idle",
        enrichment_started_at=(
            row.enrichment_started_at.isoformat() if row.enrichment_started_at else None
        ),
        enrichment_completed_at=(
            row.enrichment_completed_at.isoformat() if row.enrichment_completed_at else None
        ),
        # Gate firms_data on paid_at — pre-payment polling gets only
        # tier_report (top-5 frontend slice). Paid users see the full
        # list including Stage-2 enrichment underscore-prefixed fields.
        firms_data=(
            _public_firms_data(row.firms_data) if row.paid_at is not None else None
        ),
    )


# ----------------------------- endpoints ------------------------------

@router.post("", response_model=SearchResponse)
def create_search(
    background_tasks: BackgroundTasks,
    case_brief: str = Form(...),
    purpose: str = Form(...),
    vertical: str = Form("immigration_attorney"),
    case_id: str | None = Form(None),
    files: list[UploadFile] = File(default_factory=list),
    session: Session = Depends(get_session),
):
    """Start a new professional search.

    Multipart form (so file uploads work alongside the fields):
        case_brief: str       — user's case description (required)
        purpose: str          — short engagement label, e.g. "H-1B 2026 cap"
        vertical: str         — persona directory; defaults to immigration_attorney
        case_id: str | None   — optional existing CaseRow id to link to
        files: File[]         — up to 8 supporting docs (PDF / DOCX / TXT / MD)

    Returns the created request with status=queued. Poll GET by id for updates.
    """
    # Each search burns ~$2-3 in API + web_search calls. A meaningful
    # brief is required so agents have something to work with — 200
    # chars (~30-40 words) is the floor for useful research output.
    # Same threshold as the frontend, enforced here so the API can't be
    # bypassed via direct curl.
    MIN_BRIEF_CHARS = 200
    brief_trimmed = case_brief.strip()
    if not brief_trimmed:
        raise HTTPException(status_code=400, detail="case_brief cannot be empty")
    if len(brief_trimmed) < MIN_BRIEF_CHARS:
        raise HTTPException(
            status_code=400,
            detail=(
                f"case_brief is too short ({len(brief_trimmed)} chars). "
                f"Each search costs real compute — please describe the "
                f"specific situation, regulatory wrinkles, location, and "
                f"timeline (minimum {MIN_BRIEF_CHARS} chars)."
            ),
        )
    if not purpose.strip():
        raise HTTPException(status_code=400, detail="purpose cannot be empty")

    uploaded_notes = _build_uploaded_notes(files) if files else None

    row = ProfessionalSearchRequestRow(
        case_id=case_id,
        case_brief=case_brief.strip(),
        purpose=purpose.strip(),
        vertical=vertical,
        status="queued",
        persona_status={},
        uploaded_notes=uploaded_notes,
    )
    session.add(row)
    session.commit()
    session.refresh(row)

    # Defer the actual import until now so a missing anthropic install
    # only blows up at dispatch time, not at app startup.
    from compliance_os.web.services.professional_search_runner import run_search_sync

    background_tasks.add_task(run_search_sync, row.id)
    # Submitter is the implicit owner — return the full payload so they
    # can immediately see what they posted (handy for client retries).
    return _serialize(row, include_pii=True)


# IMPORTANT: keep static-prefix routes (`/mine/list`, `/stripe-webhook`)
# defined BEFORE the catch-all `/{request_id}` below. FastAPI matches in
# declaration order; a future regression that splits these routes across
# files could shadow the static ones if the order isn't preserved.
@router.get("/mine/list", response_model=list[SearchResponse])
def list_my_searches(
    authorization: str | None = Header(None),
    db: Session = Depends(get_session),
):
    """List the authenticated user's claimed searches.

    Path is `/mine/list` (not `/mine`) so it doesn't collide with the
    `/{request_id}` route — FastAPI's `/{x}` matches one path segment,
    so the two-segment `/mine/list` is unambiguous.
    """
    payload = get_bearer_payload(authorization, db)
    rows = (
        db.query(ProfessionalSearchRequestRow)
        .filter(ProfessionalSearchRequestRow.user_id == payload["user_id"])
        .order_by(ProfessionalSearchRequestRow.created_at.desc())
        .limit(50)
        .all()
    )
    # Authenticated owner — full PII is appropriate.
    return [_serialize(r, include_pii=True) for r in rows]


@router.get("/{request_id}", response_model=SearchResponse)
def get_search(
    request_id: str,
    authorization: str | None = Header(None),
    session: Session = Depends(get_session),
):
    """Get a search by id.

    Anonymous callers get a sanitized response (no `case_brief`, no
    `uploaded_notes`, no `stripe_customer_email`). The authenticated
    owner — and only the owner — gets the full row. We also accept any
    valid bearer for the implicit case where the user has logged in
    *between* search submission and the post-purchase claim, but match
    the user only after parsing the token.
    """
    row = session.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        raise HTTPException(status_code=404, detail="Search request not found")

    include_pii = False
    if authorization:
        try:
            payload = get_bearer_payload(authorization, session)
            if row.user_id and row.user_id == payload.get("user_id"):
                include_pii = True
        except HTTPException:
            # Bad / missing token — fall through to anonymous response.
            pass

    return _serialize(row, include_pii=include_pii)


# --- Guardian marketplace upsell ---
#
# Returns 0–2 Guardian products that fit this search's vertical+brief.
# Empty list when nothing fits (e.g. EB-5 — Guardian has no equivalent
# in-house product; the user should hire an external firm). The frontend
# only renders the upsell card when this list is non-empty, so an empty
# response is a normal valid state, not an error.


class MarketplaceMatch(BaseModel):
    sku: str
    name: str
    public_name: str | None = None
    description: str
    public_description: str | None = None
    price_cents: int
    headline: str | None = None
    public_headline: str | None = None
    cta_label: str | None = None
    public_cta_label: str | None = None
    path: str | None = None
    match_score: int
    match_reason: str


@router.get(
    "/{request_id}/marketplace-match",
    response_model=list[MarketplaceMatch],
)
def get_marketplace_match(request_id: str, session: Session = Depends(get_session)):
    """Recommend Guardian marketplace products that fit this search.

    Used by the search results page to render an upsell card above the
    external-firm tier list — closes the funnel for users whose situation
    matches a Guardian product we can fulfill in-house.
    """
    row = session.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        raise HTTPException(status_code=404, detail="Search request not found")
    from compliance_os.web.services.marketplace_match import match_products
    matches = match_products(vertical=row.vertical, case_brief=row.case_brief)
    return [MarketplaceMatch(**m) for m in matches]


# Note: there is intentionally no unauthenticated `GET /api/professional-search`
# list endpoint. Search briefs frequently contain PII (names, employers,
# regulatory specifics), and returning them across users would be a leak.
# If a list view is needed later, scope it to the authenticated user.


# ---------------------------- Report rendering ----------------------------
#
# Two output formats:
#
#   - HTML (default; opens in a new tab, prints cleanly to PDF) — the
#     polished, user-facing artifact. Firms are deduplicated across
#     personas; the same firm with rationales from multiple search
#     axes shows once with all rationales aggregated.
#
#   - Markdown (`?format=md`) — for power users who want to pipe results
#     into another tool. Keeps the persona-grouped structure since
#     markdown lacks anchor / styling affordances for cross-referencing.


PERSONA_LABELS = {
    # Generic immigration set
    "elite_boutique": "Elite boutiques",
    "startup_founder": "Startup / founder-focused",
    "litigation_contrarian": "Federal-court litigators",
    "employment_green_card": "Employment green-card counsel",
    "family_humanitarian": "Family / humanitarian counsel",
    "student_opt_status": "Student / OPT status counsel",
    # EB-5 set
    "eb5_specialist": "EB-5 specialists",
    "securities_sophisticated": "Securities-sophisticated",
    "source_of_funds": "Source-of-funds specialists",
    "i829_redeployment": "I-829 / redeployment counsel",
    "mandamus_delay": "Mandamus / delay litigators",
    "regional_center_diligence": "Regional center diligence counsel",
    # Tax attorney set
    "cross_border_tax": "Cross-border tax counsel",
    "crypto_digital_assets": "Crypto / digital-asset tax counsel",
    "disclosure_remediation": "Disclosure and remediation counsel",
    "estate_gift_cross_border": "Estate / gift cross-border counsel",
    "international_info_returns": "International information-return counsel",
    "penalty_relief": "Penalty relief and controversy counsel",
    # Corporate attorney set
    "commercial_ip_contracts": "Commercial / IP contracts counsel",
    "cross_border_equity": "Cross-border equity counsel",
    "employment_contractors": "Employment / contractor counsel",
    "m_and_a_diligence": "M&A diligence counsel",
    "safes_and_notes": "SAFE and note counsel",
    "startup_formation": "Startup formation counsel",
    # CPA set
    "audit_defense": "Audit-defense CPAs",
    "equity_crypto_tax": "Equity / crypto tax CPAs",
    "expat_tax": "Expat tax CPAs",
    "foreign_owned_entity": "Foreign-owned entity CPAs",
    "founder_accounting": "Founder-friendly accounting CPAs",
    "international_tax": "International tax CPAs",
    "state_sales_tax": "State / sales-tax CPAs",
}


def _report_vocabulary(vertical: str | None) -> dict[str, Any]:
    normalized = (vertical or "").lower()
    if (
        "attorney" in normalized
        or "lawyer" in normalized
        or normalized.startswith("immigration")
        or "eb5" in normalized
    ):
        kind = "attorney"
    elif normalized == "cpa" or "accounting" in normalized or "tax_cpa" in normalized:
        kind = "cpa"
    elif "bank" in normalized:
        kind = "bank"
    elif normalized == "caa" or "acceptance_agent" in normalized or "itin" in normalized:
        kind = "caa"
    else:
        kind = "provider"

    vocab: dict[str, Any] = {
        "is_attorney": kind == "attorney",
        "report_title": "Professional search",
        "lead_label": "Lead contact",
        "alternate_label": "Alternate contact",
        "alternate_group_label": "Same-provider alternates",
        "issue_pattern_label": "Issue pattern",
        "org_singular": "provider",
        "org_plural": "providers",
        "org_plural_title": "Providers",
        "org_dossiers": "Provider dossiers",
        "no_orgs": "No providers surfaced.",
        "retain_copy": "engaging a provider",
        "stripe_description": (
            "Full professional search report (PDF + HTML) with all provider "
            "dossiers, credentials, and verification sources."
        ),
    }
    if kind == "attorney":
        vocab.update({
            "lead_label": "Lead attorney",
            "alternate_label": "Alternate attorney",
            "alternate_group_label": "Same-firm alternates",
            "issue_pattern_label": "RFE pattern",
            "org_singular": "firm",
            "org_plural": "firms",
            "org_plural_title": "Firms",
            "org_dossiers": "Firm dossiers",
            "no_orgs": "No firms surfaced.",
            "retain_copy": "retaining a firm",
            "stripe_description": (
                "Full attorney search report (PDF + HTML) with all firm "
                "dossiers, credentials, and verification sources."
            ),
        })
    elif kind == "cpa":
        vocab.update({
            "lead_label": "Lead CPA/contact",
            "alternate_group_label": "Same-practice alternates",
            "org_singular": "CPA practice",
            "org_plural": "CPA practices",
            "org_plural_title": "CPA practices",
            "org_dossiers": "CPA practice dossiers",
            "no_orgs": "No CPA practices surfaced.",
            "retain_copy": "engaging a CPA practice",
            "stripe_description": (
                "Full CPA search report (PDF + HTML) with all practice "
                "dossiers, credentials, and verification sources."
            ),
        })
    elif kind == "bank":
        vocab.update({
            "lead_label": "Lead banker/contact",
            "alternate_group_label": "Same-provider alternates",
            "org_singular": "banking provider",
            "org_plural": "banking providers",
            "org_plural_title": "Banking providers",
            "org_dossiers": "Banking provider dossiers",
            "no_orgs": "No banking providers surfaced.",
            "retain_copy": "engaging a banking provider",
            "stripe_description": (
                "Full banking-provider search report (PDF + HTML) with all "
                "provider dossiers, credentials, and verification sources."
            ),
        })
    elif kind == "caa":
        vocab.update({
            "lead_label": "Lead acceptance agent/contact",
            "alternate_group_label": "Same-provider alternates",
            "org_singular": "acceptance agent",
            "org_plural": "acceptance agents",
            "org_plural_title": "Acceptance agents",
            "org_dossiers": "Acceptance agent dossiers",
            "no_orgs": "No acceptance agents surfaced.",
            "retain_copy": "engaging an acceptance agent",
            "stripe_description": (
                "Full acceptance-agent search report (PDF + HTML) with all "
                "provider dossiers, credentials, and verification sources."
            ),
        })
    return vocab


def _report_methodology_html(vocab: dict[str, Any]) -> str:
    if vocab["is_attorney"]:
        return (
            "<p>Firms were scored 0–100 against externally-verifiable credentials only — "
            "Chambers USA / Global rankings, AILA elected leadership (not just membership), "
            "AV Preeminent and Best Lawyers peer recognition, ABIL membership, documented "
            "PACER filings, third-party press coverage, and government-service alumni status.</p>"
            "<p>Self-published marketing — firm blog posts, &ldquo;success stories&rdquo;, "
            "self-described practice pages — was excluded from weighting. Where a firm&rsquo;s "
            "own site is cited, it is for contact information only, never as a credential source.</p>"
            "<p>Three independent search axes were dispatched in parallel: <em>elite boutiques</em> "
            "(peer-recognized specialists), <em>startup-focused</em> (firms with published positions "
            "on the specific regulatory question), and <em>federal-court litigators</em> (counsel "
            "of record in reported decisions). Convergence across axes is treated as a quality "
            "signal and surfaced via the <span class=\"cross-axis-inline\">×N</span> badge.</p>"
        )

    org_plural_title = vocab["org_plural_title"]
    org_singular = vocab["org_singular"]
    return (
        f"<p>{org_plural_title} were scored 0–100 against externally-verifiable signals "
        "relevant to the selected category: professional credentials, regulator or industry "
        "listings, third-party references, published service focus, pricing transparency, "
        "and fit with the case brief.</p>"
        f"<p>Self-published marketing is used only for {org_singular} contact and service "
        "details. It is not treated as a sole credential signal.</p>"
        "<p>Independent search axes were dispatched in parallel based on the selected category. "
        "Convergence across axes is treated as a quality signal and surfaced via the "
        "<span class=\"cross-axis-inline\">×N</span> badge.</p>"
    )


def _report_methodology_markdown(vocab: dict[str, Any]) -> str:
    if vocab["is_attorney"]:
        return (
            "Firms are scored 0–100 by externally-verifiable signals only — "
            "Chambers rankings, AILA elected leadership (not membership), "
            "AV Preeminent / Best Lawyers, ABIL membership, documented PACER "
            "filings, and third-party press. A firm's own marketing copy "
            "(blog posts, self-described practice pages) is excluded from "
            "weighting."
        )
    return (
        f"{vocab['org_plural_title']} are scored 0–100 by externally-verifiable "
        "signals relevant to the selected category: professional credentials, "
        "regulator or industry listings, third-party references, published service "
        "focus, pricing transparency, and fit with the case brief. Self-published "
        "marketing is used only for contact and service details, not as a sole "
        "credential signal."
    )


def _persona_label(persona_id: str) -> str:
    """Pretty label, with a graceful fallback for tuned personas."""
    if persona_id in PERSONA_LABELS:
        return PERSONA_LABELS[persona_id]
    if persona_id.startswith("tuned_"):
        rest = persona_id[len("tuned_"):].replace("_", " ").title()
        return f"Tuned · {rest}"
    return persona_id.replace("_", " ").title()


# Aggregation helpers live in `compliance_os.professional_search.aggregator`
# (package layer, not web layer) so the background runner can import them
# without a runner→router back-edge. Local aliases keep the existing
# call sites in this file unchanged.
from compliance_os.professional_search.aggregator import (
    aggregate_firms as _aggregate_firms_impl,
    load_persona_yamls as _load_persona_yamls_impl,
)


def _load_persona_yamls(row: ProfessionalSearchRequestRow) -> dict[str, dict]:
    return _load_persona_yamls_impl(row)


def _aggregate_firms(persona_yamls: dict[str, dict]) -> list[dict]:
    return _aggregate_firms_impl(persona_yamls)


def _score_tier(score: int | None) -> tuple[str, str]:
    """(tier-letter, color-class) for a 0-100 confidence score."""
    if score is None:
        return ("?", "tier-unknown")
    if score >= 90:
        return ("S", "tier-s")
    if score >= 80:
        return ("A", "tier-a")
    if score >= 70:
        return ("B", "tier-b")
    if score >= 60:
        return ("C", "tier-c")
    return ("D", "tier-d")


def _slug(s: str) -> str:
    out = []
    for ch in s.lower():
        if ch.isalnum():
            out.append(ch)
        elif out and out[-1] != "-":
            out.append("-")
    return "".join(out).strip("-")[:60]


def _render_html(row: ProfessionalSearchRequestRow) -> str:
    """Render a presentable, print-friendly HTML report.

    Self-contained: all CSS inlined, no external assets, opens in any
    browser. Designed for both screen viewing and Print-to-PDF.

    Data sources, in priority order:
      1. `row.firms_data` — DB-resident aggregated firm dossiers
         (canonical for any search run after this column was added)
      2. on-disk persona YAMLs — fallback for older rows
    """
    firms: list[dict] = []
    persona_yamls: dict[str, dict] = {}
    if row.firms_data:
        firms = list(row.firms_data)
        # Reconstruct persona set from firms_data so the metric tile
        # ("N search axes") still renders correctly.
        seen_personas: set[str] = set()
        for f in firms:
            for pid in f.get("_personas") or []:
                seen_personas.add(pid)
        persona_yamls = {pid: {} for pid in seen_personas}
    else:
        persona_yamls = _load_persona_yamls(row)
        firms = _aggregate_firms(persona_yamls)

    e = _html.escape
    vocab = _report_vocabulary(row.vertical)
    methodology_html = _report_methodology_html(vocab)

    started = row.created_at.strftime("%B %d, %Y") if row.created_at else "—"
    finished = row.completed_at.strftime("%B %d, %Y") if row.completed_at else None

    # Tier summary rows (top 10)
    summary_rows = []
    for f in firms[:10]:
        tier, tier_cls = _score_tier(f.get("confidence"))
        score = f.get("confidence")
        score_str = str(score) if score is not None else "—"
        loc_bits = [x for x in [f.get("city"), f.get("state")] if x]
        loc = ", ".join(loc_bits) if loc_bits else "—"
        n_axes = len(f.get("_personas") or [])
        cross = (
            f'<span class="cross-axis" title="Surfaced by {n_axes} different search axes — strongest match signal">×{n_axes}</span>'
            if n_axes >= 2
            else ""
        )
        anchor = _slug(f["name"])
        summary_rows.append(
            f'<tr class="row-{tier_cls}">'
            f'<td class="t-score"><span class="tier-pill {tier_cls}">{tier}</span><span class="num">{score_str}</span></td>'
            f'<td class="t-firm"><a href="#firm-{anchor}">{e(f["name"])}</a> {cross}</td>'
            f'<td class="t-loc">{e(loc)}</td>'
            f'<td class="t-fee">{_format_fee_range(f)}</td>'
            f"</tr>"
        )

    # Per-firm detail cards
    firm_cards: list[str] = []
    for f in firms:
        firm_cards.append(_render_firm_card(f, e, vocab))

    n_firms = len(firms)
    cross_count = sum(1 for f in firms if len(f.get("_personas") or []) >= 2)
    persona_count = len(persona_yamls)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{e(vocab["report_title"])} — {e(row.purpose)}</title>
<style>
{_REPORT_CSS}
</style>
</head>
<body>
<div class="page">
  <header class="cover">
    <div class="eyebrow">Professional search report</div>
    <h1>{e(row.purpose)}</h1>
    <div class="meta">
      <span>{e(row.vertical.replace("_", " ").title())}</span>
      <span class="dot">·</span>
      <span>{started}</span>
      {f'<span class="dot">·</span><span>finished {finished}</span>' if finished else ''}
    </div>
    <div class="metrics">
      <div class="metric"><div class="num">{n_firms}</div><div class="lbl">{e(vocab["org_plural_title"])}</div></div>
      <div class="metric"><div class="num">{persona_count}</div><div class="lbl">Search axes</div></div>
      <div class="metric"><div class="num">{cross_count}</div><div class="lbl">Cross-axis matches</div></div>
    </div>
  </header>

  <section class="block">
    <h2>Tier summary</h2>
    <p class="lede">Top {e(vocab["org_plural"])} ranked by confidence. <span class="cross-axis-inline">×N</span> means a {e(vocab["org_singular"])} was independently surfaced by N different search axes — the strongest quality signal.</p>
    <table class="tier-table">
      <thead>
        <tr><th>Tier</th><th>{e(vocab["org_singular"])}</th><th>Location</th><th>Estimated fees</th></tr>
      </thead>
      <tbody>
        {''.join(summary_rows) if summary_rows else f'<tr><td colspan="4" class="empty">{e(vocab["no_orgs"])}</td></tr>'}
      </tbody>
    </table>
  </section>

  <section class="block">
    <h2>How {e(vocab["org_plural"])} were ranked</h2>
    {methodology_html}
  </section>

  <section class="block firm-list">
    <h2>{e(vocab["org_dossiers"])}</h2>
    {''.join(firm_cards) if firm_cards else f'<p class="empty">{e(vocab["no_orgs"])}</p>'}
  </section>

  <section class="block brief">
    <h2>Case brief used for this search</h2>
    <pre>{e(row.case_brief.strip())}</pre>
  </section>

  <footer>
    <div>Generated {_dt.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")} by Guardian professional search · Report ID {e(row.id[:8])}</div>
    <div class="disclaimer">This report is research output, not legal or tax advice. Verify each credential against its cited source before {e(vocab["retain_copy"])}.</div>
  </footer>
</div>
</body>
</html>
"""
    return html


def _format_fee_range(f: dict) -> str:
    lo = f.get("petition_fee_low")
    hi = f.get("petition_fee_high") or lo
    if lo is None:
        return '<span class="fees-missing-inline">quote required</span>'
    if hi is None or hi == lo:
        return f"${int(lo):,}"
    return f"${int(lo):,}–${int(hi):,}"


def _report_text(value: Any) -> str | None:
    if value is None:
        return None
    if isinstance(value, str):
        text = value.strip()
        return text or None
    if isinstance(value, (int, float, bool)):
        return str(value)
    if isinstance(value, (list, tuple)):
        parts = [_report_text(item) for item in value]
        text = "; ".join(part for part in parts if part)
        return text or None
    return None


def _report_text_list(value: Any) -> list[str]:
    if isinstance(value, (list, tuple)):
        return [text for item in value if (text := _report_text(item))]
    text = _report_text(value)
    return [text] if text else []


def _report_sequence(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, (list, tuple)):
        return list(value)
    return [value]


def _band_label(source: Any, band: Any, year: Any) -> str | None:
    band_text = _report_text(band)
    if not band_text:
        return None
    if not band_text.lower().startswith("band"):
        band_text = f"Band {band_text}"
    source_text = _report_text(source)
    year_text = _report_text(year)
    label = f"{source_text} {band_text}" if source_text else band_text
    if year_text:
        label = f"{label} ({year_text})"
    return label


def _render_source_item(source: Any, e) -> str | None:
    if isinstance(source, dict):
        url = _report_text(source.get("url"))
        title = _report_text(source.get("title"))
        note = _report_text(source.get("note"))
        verified = source.get("verified")
        status = (
            "verified" if verified is True
            else "unverified" if verified is False
            else None
        )
        details = " - ".join(part for part in [title, status, note] if part)
    else:
        url = _report_text(source)
        details = None
    if not url:
        return None
    label = details or url
    if url.startswith(("http://", "https://")):
        return (
            f'<li><a href="{e(url)}" target="_blank" rel="noopener">'
            f"{e(label)}</a></li>"
        )
    return f"<li>{e(label)}</li>"


def _render_stage_two_enrichment(f: dict, e, vocab: dict[str, Any]) -> str:
    warning = _report_text(f.get("_individual_vs_firm_band_gap_warning"))
    lead_band = _band_label(
        f.get("_lead_attorney_band_source"),
        f.get("_lead_attorney_band"),
        f.get("_lead_attorney_band_year"),
    )
    practice_focus = _report_text(f.get("_lead_attorney_practice_focus"))
    lead_creds = _report_text_list(f.get("_lead_attorney_credentials"))
    rfe_pattern = _report_text(f.get("_rfe_pattern"))

    alternates = [
        item for item in _report_sequence(f.get("_alternate_attorneys"))
        if isinstance(item, dict) and _report_text(item.get("name"))
    ]
    verified_sources = [
        item for item in _report_sequence(f.get("_verified_sources"))
        if _render_source_item(item, e)
    ]

    if not any([
        warning,
        lead_band,
        practice_focus,
        lead_creds,
        rfe_pattern,
        alternates,
        verified_sources,
    ]):
        return ""

    lead_name = (
        _report_text(f.get("lead_attorney") or f.get("lead_contact"))
        or vocab["lead_label"]
    )
    lead_consults = f.get("_lead_attorney_takes_outside_consults")
    consult_note = ""
    if lead_consults is True:
        consult_note = '<span class="consult-note">Reported as taking outside consults</span>'
    elif lead_consults is False:
        consult_note = (
            '<span class="consult-note consult-limited">'
            "May not take outside consults</span>"
        )

    lead_parts = [f"<strong>{e(lead_name)}</strong>"]
    if lead_band:
        lead_parts.append(f'<span class="band-pill">{e(lead_band)}</span>')
    if consult_note:
        lead_parts.append(consult_note)

    warning_html = (
        f'<div class="enrichment-warning">{e(warning)}</div>' if warning else ""
    )
    practice_html = (
        '<div class="enrichment-row"><span>Practice focus</span>'
        f"<p>{e(practice_focus)}</p></div>"
        if practice_focus else ""
    )
    creds_html = (
        "<div class='enrichment-row'><span>Individual credentials</span>"
        "<ul class='creds'>"
        + "".join(f"<li>{e(cred)}</li>" for cred in lead_creds)
        + "</ul></div>"
        if lead_creds else ""
    )
    rfe_html = (
        f'<div class="enrichment-row"><span>{e(vocab["issue_pattern_label"])}</span>'
        f"<p>{e(rfe_pattern)}</p></div>"
        if rfe_pattern else ""
    )

    alternate_html = ""
    if alternates:
        items: list[str] = []
        for alternate in alternates[:3]:
            alt_name = _report_text(alternate.get("name")) or vocab["alternate_label"]
            alt_band = _band_label(
                alternate.get("band_source"),
                alternate.get("band"),
                alternate.get("band_year"),
            )
            alt_fit = _report_text(alternate.get("fit_for_case"))
            alt_focus = _report_text(alternate.get("practice_focus"))
            alt_consults = alternate.get("takes_outside_consults")
            alt_consult_note = ""
            if alt_consults is False:
                alt_consult_note = (
                    '<span class="consult-note consult-limited">'
                    "retained matters only</span>"
                )
            elif alt_consults is True:
                alt_consult_note = '<span class="consult-note">takes outside consults</span>'
            band_html = (
                f'<span class="band-pill alt-band">{e(alt_band)}</span>'
                if alt_band else ""
            )
            detail = alt_fit or alt_focus
            detail_html = f"<p>{e(detail)}</p>" if detail else ""
            items.append(
                "<li>"
                f"<div><strong>{e(alt_name)}</strong> "
                f"{band_html} {alt_consult_note}</div>"
                f"{detail_html}"
                "</li>"
            )
        alternate_html = (
            f"<div class='enrichment-row'><span>{e(vocab['alternate_group_label'])}</span>"
            "<ul class='alt-list'>"
            + "".join(items)
            + "</ul></div>"
        )

    verified_sources_html = ""
    if verified_sources:
        items = [_render_source_item(source, e) for source in verified_sources]
        verified_sources_html = (
            "<div class='enrichment-row'><span>Individual verification sources</span>"
            "<ul class='sources verified-sources'>"
            + "".join(item for item in items if item)
            + "</ul></div>"
        )

    return f"""
  <section class="enrichment">
    <h4>Stage 2 individual verification</h4>
    {warning_html}
    <div class="lead-line">{" ".join(lead_parts)}</div>
    {practice_html}
    {creds_html}
    {alternate_html}
    {rfe_html}
    {verified_sources_html}
  </section>
"""


def _render_firm_card(f: dict, e, vocab: dict[str, Any]) -> str:
    name = f.get("name", "(unnamed)")
    anchor = _slug(name)
    score = f.get("confidence")
    tier, tier_cls = _score_tier(score)
    score_str = str(score) if score is not None else "—"

    contact_bits: list[str] = []
    lead_contact = f.get("lead_attorney") or f.get("lead_contact")
    if lead_contact:
        role = f.get("role")
        contact_bits.append(
            f"<strong>{e(lead_contact)}</strong>"
            + (f" <span class='role'>({e(role)})</span>" if role else "")
        )
    loc = ", ".join([x for x in [f.get("city"), f.get("state")] if x])
    if loc:
        contact_bits.append(e(loc))
    if f.get("phone"):
        contact_bits.append(e(f["phone"]))
    if f.get("email"):
        contact_bits.append(f'<a href="mailto:{e(f["email"])}">{e(f["email"])}</a>')
    if f.get("website"):
        contact_bits.append(f'<a href="{e(f["website"])}" target="_blank" rel="noopener">{e(f["website"].replace("https://", "").replace("http://", "").rstrip("/"))}</a>')

    fee_html = ""
    fee_rows: list[str] = []
    if f.get("consultation_fee_low") is not None:
        lo = f["consultation_fee_low"]
        hi = f.get("consultation_fee_high") or lo
        amt = f"${int(lo):,}" + (f"–${int(hi):,}" if hi != lo else "")
        fee_rows.append(f'<dt>Consultation</dt><dd>{amt}</dd>')
    if f.get("petition_fee_low") is not None:
        lo = f["petition_fee_low"]
        hi = f.get("petition_fee_high") or lo
        amt = f"${int(lo):,}" + (f"–${int(hi):,}" if hi != lo else "")
        label = f.get("service_fee_label", "Primary engagement")
        fee_rows.append(f'<dt>{e(label)}</dt><dd>{amt}</dd>')
    if not fee_rows:
        fee_rows.append(
            '<dt>Fees</dt><dd class="fees-missing">'
            'Not publicly disclosed — request quote at consultation</dd>'
        )
    basis = f.get("fee_basis")
    basis_html = (
        f'<div class="fee-basis">Source: {e(basis)}</div>' if basis else ""
    )
    fee_html = f'<dl class="fees">{"".join(fee_rows)}</dl>{basis_html}'

    # `_why_fits` shape evolution: the current aggregator writes
    # `[persona_id, text]` 2-lists, but older rows in firms_data have
    # bare strings (the why_fit text alone, no persona tag). Normalize
    # to (pid_or_None, text) tuples so the renderer survives both.
    raw_why = f.get("_why_fits") or []
    why_fits: list[tuple[str | None, str]] = []
    for item in raw_why:
        if isinstance(item, (list, tuple)):
            if len(item) >= 2:
                why_fits.append((item[0], item[1]))
            elif len(item) == 1:
                why_fits.append((None, item[0]))
        elif isinstance(item, str):
            why_fits.append((None, item))
        elif isinstance(item, dict):
            why_fits.append((item.get("persona_id"), item.get("text") or item.get("why_fit") or ""))

    why_html = ""
    if len(why_fits) == 1:
        why_html = f"<p>{e(why_fits[0][1])}</p>"
    elif len(why_fits) > 1:
        parts = []
        for pid, text in why_fits:
            if pid:
                parts.append(
                    f"<div class='persona-take'><div class='persona-tag'>{e(_persona_label(pid))}</div><p>{e(text)}</p></div>"
                )
            else:
                parts.append(f"<div class='persona-take'><p>{e(text)}</p></div>")
        why_html = "".join(parts)

    creds_html = ""
    if f.get("_credentials"):
        creds_html = (
            "<h4>Credentials <span class='note'>(externally verifiable)</span></h4>"
            "<ul class='creds'>"
            + "".join(f"<li>{e(c)}</li>" for c in f["_credentials"])
            + "</ul>"
        )

    lit_html = ""
    if f.get("litigation_capability"):
        lit_html = (
            f"<h4>Litigation capability</h4><p class='lit'>{e(f['litigation_capability'])}</p>"
        )

    risks_html = ""
    if f.get("_risks"):
        risks_html = (
            "<h4>Caveats</h4><ul class='risks'>"
            + "".join(f"<li>{e(r)}</li>" for r in f["_risks"])
            + "</ul>"
        )

    sources_html = ""
    if f.get("_sources"):
        sources_html = (
            "<h4>Verification sources</h4><ul class='sources'>"
            + "".join(
                f'<li><a href="{e(s)}" target="_blank" rel="noopener">{e(s)}</a></li>'
                for s in f["_sources"]
            )
            + "</ul>"
        )

    stage_two_html = _render_stage_two_enrichment(f, e, vocab)

    persona_pills = " ".join(
        f'<span class="persona-pill">{e(_persona_label(p))}</span>'
        for p in (f.get("_personas") or [])
    )
    n_axes = len(f.get("_personas") or [])
    cross_html = (
        f'<span class="cross-axis-card" title="Surfaced by {n_axes} different search axes">×{n_axes} cross-axis match</span>'
        if n_axes >= 2 else ""
    )

    return f"""
<article class="firm" id="firm-{anchor}">
  <div class="firm-head">
    <div class="firm-score">
      <span class="tier-pill {tier_cls}">{tier}</span>
      <span class="firm-score-num">{score_str}</span>
    </div>
    <div class="firm-title">
      <h3>{e(name)}</h3>
      <div class="firm-axes">{persona_pills} {cross_html}</div>
    </div>
  </div>
  <div class="firm-contact">{" · ".join(contact_bits)}</div>
  {fee_html}
  {stage_two_html}
  {why_html}
  {creds_html}
  {lit_html}
  {risks_html}
  {sources_html}
</article>
"""


_REPORT_CSS = """
:root {
  --ink: #0d1424;
  --muted: #556480;
  --muted-2: #7b8ba5;
  --hairline: #e4edf7;
  --soft: #f5f7fb;
  --card: #ffffff;
  --accent: #2f5bae;
  --accent-soft: #eaf2ff;
}
* { box-sizing: border-box; }
html { -webkit-text-size-adjust: 100%; }
body {
  margin: 0;
  font-family: -apple-system, BlinkMacSystemFont, "Inter", "Helvetica Neue", Arial, sans-serif;
  font-size: 15px;
  line-height: 1.55;
  color: var(--ink);
  background: linear-gradient(180deg, #edf3f9 0%, #f4f7fb 60%, #fbfcfe 100%);
  -webkit-font-smoothing: antialiased;
}
.page {
  max-width: 880px;
  margin: 48px auto;
  padding: 0 32px 64px;
}
header.cover {
  padding: 56px 0 40px;
  border-bottom: 1px solid var(--hairline);
}
.eyebrow {
  font-size: 11px;
  font-weight: 700;
  letter-spacing: 0.22em;
  text-transform: uppercase;
  color: var(--muted-2);
}
header.cover h1 {
  margin: 14px 0 14px;
  font-family: "Charter", "Georgia", "Times New Roman", serif;
  font-size: 44px;
  font-weight: 700;
  line-height: 1.1;
  letter-spacing: -0.01em;
  color: var(--ink);
}
.meta { color: var(--muted); font-size: 14px; }
.meta .dot { color: #c4cad6; margin: 0 8px; }
.metrics {
  display: flex;
  gap: 12px;
  margin-top: 28px;
}
.metric {
  flex: 1;
  background: var(--card);
  border: 1px solid var(--hairline);
  border-radius: 14px;
  padding: 16px 20px;
  box-shadow: 0 6px 20px rgba(56,85,131,0.04);
}
.metric .num {
  font-size: 28px;
  font-weight: 700;
  color: var(--ink);
  line-height: 1;
}
.metric .lbl {
  font-size: 11px;
  font-weight: 600;
  letter-spacing: 0.16em;
  text-transform: uppercase;
  color: var(--muted-2);
  margin-top: 8px;
}
.block { margin-top: 56px; }
.block h2 {
  font-family: "Charter", "Georgia", "Times New Roman", serif;
  font-size: 26px;
  font-weight: 700;
  margin: 0 0 6px;
  color: var(--ink);
}
.block .lede { color: var(--muted); margin: 0 0 20px; font-size: 14px; }
.block p { color: var(--muted); margin: 0 0 14px; max-width: 64ch; }
.block p strong, .block p em { color: var(--ink); font-style: normal; font-weight: 600; }

/* tier table */
.tier-table {
  width: 100%;
  border-collapse: separate;
  border-spacing: 0;
  background: var(--card);
  border: 1px solid var(--hairline);
  border-radius: 14px;
  overflow: hidden;
  box-shadow: 0 6px 20px rgba(56,85,131,0.04);
}
.tier-table th {
  text-align: left;
  font-size: 11px;
  font-weight: 700;
  letter-spacing: 0.14em;
  text-transform: uppercase;
  color: var(--muted-2);
  background: var(--soft);
  padding: 12px 16px;
  border-bottom: 1px solid var(--hairline);
}
.tier-table td {
  padding: 14px 16px;
  border-bottom: 1px solid var(--hairline);
  vertical-align: middle;
  font-size: 14px;
}
.tier-table tr:last-child td { border-bottom: 0; }
.tier-table .t-score { width: 100px; }
.tier-table .t-score .num { margin-left: 8px; color: var(--muted); font-variant-numeric: tabular-nums; font-size: 13px; }
.tier-table .t-firm a { color: var(--ink); font-weight: 600; text-decoration: none; }
.tier-table .t-firm a:hover { color: var(--accent); }
.tier-table .t-loc { color: var(--muted); }
.tier-table .t-fee { color: var(--muted); font-variant-numeric: tabular-nums; white-space: nowrap; }
.tier-table .empty { color: var(--muted-2); text-align: center; padding: 24px; }

/* tier pills */
.tier-pill {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  width: 28px;
  height: 28px;
  border-radius: 8px;
  font-weight: 700;
  font-size: 13px;
  letter-spacing: 0;
  color: white;
}
.tier-pill.tier-s { background: #1f6f43; }
.tier-pill.tier-a { background: #2f5bae; }
.tier-pill.tier-b { background: #5b8dee; }
.tier-pill.tier-c { background: #9ba8c4; }
.tier-pill.tier-d { background: #c4cad6; color: #5a6478; }
.tier-pill.tier-unknown { background: #e4edf7; color: #7b8ba5; }

.cross-axis {
  display: inline-flex;
  align-items: center;
  margin-left: 8px;
  padding: 2px 8px;
  font-size: 11px;
  font-weight: 700;
  letter-spacing: 0.04em;
  background: #fff6ea;
  color: #9c5a1c;
  border: 1px solid #ffe3c9;
  border-radius: 999px;
  vertical-align: middle;
}
.cross-axis-inline {
  display: inline-block;
  padding: 1px 6px;
  font-size: 11px;
  font-weight: 700;
  background: #fff6ea;
  color: #9c5a1c;
  border: 1px solid #ffe3c9;
  border-radius: 999px;
}

/* firm cards */
.firm {
  background: var(--card);
  border: 1px solid var(--hairline);
  border-radius: 16px;
  padding: 28px 32px;
  margin-top: 16px;
  box-shadow: 0 6px 20px rgba(56,85,131,0.04);
}
.firm-head {
  display: flex;
  gap: 18px;
  align-items: flex-start;
}
.firm-score {
  display: flex;
  flex-direction: column;
  align-items: center;
  gap: 4px;
  padding-top: 4px;
}
.firm-score .tier-pill {
  width: 40px;
  height: 40px;
  font-size: 17px;
  border-radius: 12px;
}
.firm-score-num {
  font-size: 13px;
  color: var(--muted-2);
  font-variant-numeric: tabular-nums;
  font-weight: 600;
}
.firm-title { flex: 1; min-width: 0; }
.firm-title h3 {
  margin: 0 0 6px;
  font-family: "Charter", "Georgia", "Times New Roman", serif;
  font-size: 22px;
  font-weight: 700;
  color: var(--ink);
}
.firm-axes { display: flex; flex-wrap: wrap; gap: 6px; align-items: center; }
.persona-pill {
  display: inline-block;
  font-size: 10.5px;
  font-weight: 600;
  letter-spacing: 0.06em;
  text-transform: uppercase;
  color: var(--muted);
  background: var(--soft);
  border: 1px solid var(--hairline);
  border-radius: 999px;
  padding: 3px 9px;
}
.cross-axis-card {
  display: inline-block;
  font-size: 10.5px;
  font-weight: 700;
  letter-spacing: 0.06em;
  text-transform: uppercase;
  color: #9c5a1c;
  background: #fff6ea;
  border: 1px solid #ffe3c9;
  border-radius: 999px;
  padding: 3px 9px;
}

.firm-contact {
  margin-top: 14px;
  padding: 12px 14px;
  background: var(--soft);
  border-radius: 10px;
  font-size: 13.5px;
  color: var(--muted);
  line-height: 1.6;
}
.firm-contact strong { color: var(--ink); }
.firm-contact .role { color: var(--muted-2); font-weight: 400; }
.firm-contact a { color: var(--accent); text-decoration: none; }
.firm-contact a:hover { text-decoration: underline; }

.firm h4 {
  font-size: 11px;
  font-weight: 700;
  letter-spacing: 0.16em;
  text-transform: uppercase;
  color: var(--muted-2);
  margin: 22px 0 10px;
}
.firm h4 .note { font-weight: 500; text-transform: none; letter-spacing: 0; color: var(--muted-2); }
.firm p { color: var(--ink); font-size: 14.5px; margin: 0 0 12px; line-height: 1.65; }

.persona-take {
  position: relative;
  padding-left: 16px;
  margin-bottom: 16px;
}
.persona-take::before {
  content: '';
  position: absolute;
  left: 0; top: 4px; bottom: 4px;
  width: 3px;
  background: var(--accent);
  border-radius: 2px;
  opacity: 0.55;
}
.persona-take .persona-tag {
  font-size: 10.5px;
  font-weight: 700;
  letter-spacing: 0.16em;
  text-transform: uppercase;
  color: var(--accent);
  margin-bottom: 4px;
}

dl.fees {
  display: grid;
  grid-template-columns: max-content 1fr;
  gap: 6px 18px;
  margin: 16px 0 4px;
  padding: 14px 18px;
  background: var(--soft);
  border-radius: 10px;
  font-size: 13.5px;
}
dl.fees dt { color: var(--muted-2); font-weight: 600; }
dl.fees dd { margin: 0; color: var(--ink); font-variant-numeric: tabular-nums; }
dl.fees dd.fees-missing { color: var(--muted); font-style: italic; font-variant-numeric: normal; }
.fee-basis {
  margin-top: 4px;
  font-size: 11px;
  color: var(--muted-2);
  font-style: italic;
}
.fees-missing-inline { color: var(--muted-2); font-style: italic; font-size: 12.5px; }

.enrichment {
  margin-top: 18px;
  padding: 16px 18px;
  background: #f7fbf8;
  border: 1px solid #dcefe1;
  border-radius: 12px;
  break-inside: avoid;
}
.enrichment h4 {
  margin-top: 0;
  color: #2f6f43;
}
.lead-line {
  display: flex;
  align-items: center;
  gap: 8px;
  flex-wrap: wrap;
  color: var(--ink);
  font-size: 14px;
  margin-bottom: 10px;
}
.band-pill {
  display: inline-block;
  padding: 2px 8px;
  border-radius: 999px;
  background: #eaf6ed;
  border: 1px solid #cbe6d3;
  color: #2f6f43;
  font-size: 11px;
  font-weight: 700;
  letter-spacing: 0.02em;
}
.consult-note {
  color: #2f6f43;
  font-size: 12px;
  font-weight: 600;
}
.consult-note.consult-limited {
  color: #9c5a1c;
}
.enrichment-warning {
  margin: 0 0 12px;
  padding: 9px 11px;
  border-radius: 9px;
  background: #fff6ea;
  border: 1px solid #ffe3c9;
  color: #7c4415;
  font-size: 13px;
  line-height: 1.5;
}
.enrichment-row {
  display: grid;
  grid-template-columns: 136px minmax(0, 1fr);
  gap: 12px;
  margin-top: 10px;
}
.enrichment-row > span {
  color: var(--muted-2);
  font-size: 11px;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
}
.enrichment-row p {
  margin: 0;
  font-size: 13.5px;
  color: var(--ink);
}
.alt-list {
  list-style: none;
  padding: 0;
  margin: 0;
  display: grid;
  gap: 8px;
}
.alt-list li {
  padding: 10px 12px;
  background: white;
  border: 1px solid #e4edf7;
  border-radius: 10px;
}
.alt-list p {
  margin-top: 4px;
  color: var(--muted);
}
.verified-sources {
  margin-top: 2px;
}

ul.creds, ul.risks, ul.sources {
  list-style: none;
  padding: 0;
  margin: 0;
  display: grid;
  gap: 4px;
}
ul.creds li, ul.risks li, ul.sources li {
  position: relative;
  padding-left: 18px;
  font-size: 13.5px;
  color: var(--ink);
}
ul.creds li::before {
  content: '✓';
  position: absolute;
  left: 0; top: 0;
  color: #1f6f43;
  font-weight: 700;
}
ul.risks li {
  color: var(--muted);
}
ul.risks li::before {
  content: '!';
  position: absolute;
  left: 4px; top: 0;
  color: #9c5a1c;
  font-weight: 700;
}
ul.sources li::before {
  content: '↗';
  position: absolute;
  left: 0; top: 0;
  color: var(--muted-2);
}
ul.sources a {
  color: var(--accent);
  text-decoration: none;
  font-size: 12.5px;
  word-break: break-all;
}
ul.sources a:hover { text-decoration: underline; }
.firm p.lit { color: var(--muted); font-size: 13.5px; margin-bottom: 4px; }

.brief pre {
  background: var(--soft);
  border: 1px solid var(--hairline);
  border-radius: 12px;
  padding: 18px 22px;
  white-space: pre-wrap;
  word-wrap: break-word;
  font-family: "SF Mono", "Menlo", "Consolas", monospace;
  font-size: 12.5px;
  line-height: 1.65;
  color: var(--muted);
}

footer {
  margin-top: 64px;
  padding-top: 24px;
  border-top: 1px solid var(--hairline);
  font-size: 12px;
  color: var(--muted-2);
}
footer .disclaimer { margin-top: 6px; font-style: italic; }

@media (max-width: 640px) {
  .page { padding: 0 18px 48px; margin: 24px auto; }
  header.cover h1 { font-size: 32px; }
  .metrics { flex-direction: column; }
  .firm { padding: 22px; }
  .firm-head { flex-direction: row; }
  .enrichment-row { grid-template-columns: 1fr; gap: 4px; }
}

@media print {
  body { background: white; }
  .page { margin: 0; max-width: none; padding: 0 24px; }
  header.cover { padding: 24px 0; }
  .firm, .metric, .tier-table, .enrichment { box-shadow: none; break-inside: avoid; }
  .block { margin-top: 32px; }
  a { color: var(--ink); }
  ul.sources a { color: var(--muted); }
}
"""


def _render_markdown(row: ProfessionalSearchRequestRow) -> str:
    """Render a portable markdown report of every firm surfaced.

    Reads the per-persona YAMLs that the runner saved (which carry the
    rich data — credentials, why_fit, sources, fees) and emits sections
    grouped by persona. Firms appearing in multiple personas are listed
    under each (with a note) — this is intentional, since the cross-axis
    overlap is itself a quality signal.
    """
    sections: list[str] = []
    vocab = _report_vocabulary(row.vertical)

    sections.append(f"# {vocab['report_title']} — {row.purpose}")
    sections.append("")
    sections.append(f"**Vertical:** {row.vertical.replace('_', ' ')}  ")
    sections.append(f"**Started:** {row.created_at.isoformat() if row.created_at else '-'}  ")
    if row.completed_at:
        sections.append(f"**Finished:** {row.completed_at.isoformat()}  ")
    sections.append(f"**Status:** {row.status}")
    sections.append("")

    sections.append("## Case brief")
    sections.append("")
    sections.append("```")
    sections.append(row.case_brief.strip())
    sections.append("```")
    sections.append("")

    sections.append("## Methodology")
    sections.append("")
    sections.append(_report_methodology_markdown(vocab))
    sections.append("")

    persona_status = row.persona_status or {}
    any_firms = False

    for persona_id, status in persona_status.items():
        title = persona_id.replace("_", " ").title()
        sections.append(f"## {title}")
        sections.append("")

        if status.get("status") != "complete":
            sections.append("_This search axis did not produce a usable result._")
            sections.append("")
            continue

        path_str = status.get("output_path")
        if not path_str:
            sections.append("_Saved result unavailable for this search axis._")
            sections.append("")
            continue
        path = Path(path_str)
        if not path.exists():
            sections.append("_Saved result unavailable for this search axis._")
            sections.append("")
            continue

        try:
            doc = _yaml.safe_load(path.read_text()) or {}
        except Exception:
            sections.append("_Saved result unavailable for this search axis._")
            sections.append("")
            continue

        firms = doc.get("firms") or []
        if not firms:
            sections.append(f"_No {vocab['org_plural']} returned._")
            sections.append("")
            continue

        any_firms = True
        firms_sorted = sorted(firms, key=lambda f: -(f.get("confidence") or 0))

        for firm in firms_sorted:
            name = firm.get("name", "(unnamed)")
            score = firm.get("confidence")
            score_str = f" — score {score}" if score is not None else ""
            sections.append(f"### {name}{score_str}")
            sections.append("")

            # Contact line
            contact_bits = []
            lead_contact = firm.get("lead_attorney") or firm.get("lead_contact")
            if lead_contact:
                role = firm.get("role")
                contact_bits.append(
                    f"**{lead_contact}**" + (f" ({role})" if role else "")
                )
            loc = ", ".join([x for x in [firm.get("city"), firm.get("state")] if x])
            if loc:
                contact_bits.append(loc)
            if firm.get("phone"):
                contact_bits.append(firm["phone"])
            if firm.get("email"):
                contact_bits.append(firm["email"])
            if firm.get("website"):
                contact_bits.append(firm["website"])
            if contact_bits:
                sections.append(" · ".join(contact_bits))
                sections.append("")

            # Fees
            fee_lines = []
            if firm.get("consultation_fee_low") is not None:
                lo = firm["consultation_fee_low"]
                hi = firm.get("consultation_fee_high") or lo
                fee_lines.append(
                    f"- Consultation: ${lo:,.0f}"
                    + (f"–${hi:,.0f}" if hi != lo else "")
                )
            if firm.get("petition_fee_low") is not None:
                lo = firm["petition_fee_low"]
                hi = firm.get("petition_fee_high") or lo
                label = firm.get("service_fee_label", "Primary service")
                fee_lines.append(
                    f"- {label}: ${lo:,.0f}"
                    + (f"–${hi:,.0f}" if hi != lo else "")
                )
            if fee_lines:
                sections.append("**Fees**")
                sections.extend(fee_lines)
                sections.append("")

            if firm.get("why_fit"):
                sections.append(f"**Why this {vocab['org_singular']}**")
                sections.append("")
                sections.append(firm["why_fit"].strip())
                sections.append("")

            creds = firm.get("credentials") or []
            if creds:
                sections.append("**Credentials (externally verifiable)**")
                for c in creds:
                    sections.append(f"- {c}")
                sections.append("")

            if firm.get("litigation_capability"):
                sections.append("**Litigation capability**")
                sections.append("")
                sections.append(firm["litigation_capability"].strip())
                sections.append("")

            risks = firm.get("risks") or []
            if risks:
                sections.append("**Risks / caveats**")
                for r in risks:
                    sections.append(f"- {r}")
                sections.append("")

            sources = firm.get("sources") or []
            if sources:
                sections.append("**Sources**")
                for s in sources:
                    sections.append(f"- <{s}>")
                sections.append("")

            sections.append("---")
            sections.append("")

    if not any_firms:
        sections.append(
            f"_No {vocab['org_plural']} were ingested. "
            "Please rerun the search or contact support._"
        )
        sections.append("")

    sections.append("")
    sections.append("_Generated by compliance-os professional search._")
    return "\n".join(sections)


@router.get("/{request_id}/download")
def download_search(
    request_id: str,
    format: str = "html",
    session: Session = Depends(get_session),
):
    """Render the full search result as a polished report.

    Two formats:
      - `html` (default): self-contained styled page, opens inline in
        the browser, prints cleanly to PDF. Firms are deduplicated
        across personas with rationales aggregated — one card per
        unique firm, cross-axis matches called out.
      - `md`: portable markdown, served as an attachment. Keeps the
        per-persona grouping for users who want to see each search
        axis separately.
    """
    row = session.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        raise HTTPException(status_code=404, detail="Search request not found")
    if row.status != "complete":
        raise HTTPException(
            status_code=409,
            detail=f"Search not complete (status={row.status}); cannot download",
        )
    if row.paid_at is None:
        raise HTTPException(
            status_code=402,  # Payment Required — semantically correct
            detail="This report is locked. Unlock it for $15 via the checkout endpoint.",
        )

    safe_purpose = "".join(
        c if c.isalnum() or c in "-_" else "-" for c in row.purpose
    ).strip("-")[:60] or "report"
    filename_prefix = "professional-search"

    # Bail early when there's nothing to render — applies to ALL formats
    # (md / html / pdf) so a stale row can't return a polished-but-empty
    # report in any guise. Must come before any format-specific branch.
    if not row.firms_data and not _load_persona_yamls(row):
        raise HTTPException(
            status_code=410,
            detail=(
                "This search's underlying result data is no longer available "
                "(YAMLs missing and no DB-resident dossiers). Run a fresh "
                "search to regenerate the report."
            ),
        )

    if format == "md":
        body = _render_markdown(row)
        filename = f"{filename_prefix}-{safe_purpose}-{request_id[:8]}.md"
        return Response(
            content=body,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if format == "pdf":
        # Lazy import — WeasyPrint pulls in cairo/pango at import time.
        try:
            from weasyprint import HTML as _WeasyHTML
        except ImportError:
            raise HTTPException(
                status_code=503,
                detail="PDF rendering not available — server missing weasyprint",
            )
        html_body = _render_html(row)
        pdf_bytes = _WeasyHTML(string=html_body).write_pdf()
        filename = f"{filename_prefix}-{safe_purpose}-{request_id[:8]}.pdf"
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # Default: polished HTML, inline so it opens as a page (not a download).
    body = _render_html(row)
    filename = f"{filename_prefix}-{safe_purpose}-{request_id[:8]}.html"
    return Response(
        content=body,
        media_type="text/html; charset=utf-8",
        headers={"Content-Disposition": f'inline; filename="{filename}"'},
    )


# ---------------------------- Stripe paywall ----------------------------
#
# Pay-first model: anyone can pay (no account required). Stripe captures
# the email at checkout. Post-purchase, the user is sent to a success
# page with a "save this to your account" CTA — signing up there calls
# /claim, which links the search row to the new (or existing) user.


def _stripe():
    """Lazy-import + configure the Stripe SDK so the app boots without keys."""
    if not settings.stripe_secret_key:
        raise HTTPException(
            status_code=503,
            detail="Stripe not configured — STRIPE_SECRET_KEY is not set",
        )
    import stripe as _s

    _s.api_key = settings.stripe_secret_key
    return _s


@router.post("/{request_id}/checkout")
def create_checkout_session(
    request_id: str,
    background_tasks: BackgroundTasks,
    authorization: str | None = Header(None),
    session: Session = Depends(get_session),
):
    """Create a Stripe Checkout Session for a completed-but-unpaid search.

    Returns a JSON `{ url }` the client redirects to. The Checkout Session's
    `success_url` redirects back to `/find-lawyer/{id}/paid?session_id=...`,
    where the frontend polls until the webhook fires and `paid_at` is set.

    Pro free-pass: when the caller is an authenticated paid-Pro user and
    has not yet consumed their 1-search-per-period grant, we skip Stripe
    entirely — mark the row paid with `pro_free_grant_at` set, and return
    a direct redirect to the /paid page.
    """
    row = session.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        raise HTTPException(status_code=404, detail="Search request not found")
    if row.status != "complete":
        raise HTTPException(
            status_code=409, detail=f"Search not complete (status={row.status})"
        )
    if row.paid_at is not None:
        return {"url": f"{settings.public_app_url}/find-lawyer/{request_id}/paid"}

    # Pro free-search short-circuit. Anonymous callers fall through to the
    # paid Stripe flow below — quietly, no error. We only consult auth
    # when a token is present so this endpoint stays open to logged-out
    # users (the dominant pre-signup flow).
    if authorization:
        try:
            payload = get_bearer_payload(authorization, session)
            user = session.query(UserRow).filter(UserRow.id == payload["user_id"]).first()
        except HTTPException:
            user = None  # bad/expired token → just fall through to paid path
        if user is not None:
            from compliance_os.web.services.subscription_service import (
                get_pro_search_quota,
            )

            # Includes pro_trial users now — see subscription_service
            # `get_pro_search_quota` for the entitlement rule.
            pro_quota = get_pro_search_quota(user, session)
            if pro_quota.has_free_search:
                    now = _dt.datetime.utcnow()
                    row.paid_at = now
                    row.pro_free_grant_at = now
                    row.user_id = user.id
                    if not row.stripe_customer_email:
                        row.stripe_customer_email = user.email
                    session.commit()
                    logger.info(
                        "pro_free_grant: user %s consumed free search for %s",
                        user.id,
                        request_id,
                    )
                    return {
                        "url": f"{settings.public_app_url}/find-lawyer/{request_id}/paid",
                        "pro_free_grant": True,
                    }

    stripe = _stripe()
    vocab = _report_vocabulary(row.vertical)
    try:
        checkout = stripe.checkout.Session.create(
            mode="payment",
            payment_method_types=["card"],
            # Force a Stripe Customer for every $15 purchase so we can
            # later create a Pro trial subscription that bills the saved
            # card without re-prompting (see /start-pro-trial below).
            customer_creation="always",
            # Save the payment method to the customer so off-session
            # subscription creation can charge it after the trial ends.
            payment_intent_data={"setup_future_usage": "off_session"},
            line_items=[
                {
                    "price_data": {
                        "currency": "usd",
                        "product_data": {
                            "name": f"Guardian — {row.purpose}",
                            "description": vocab["stripe_description"],
                        },
                        "unit_amount": settings.stripe_report_price_cents,
                    },
                    "quantity": 1,
                }
            ],
            # The session_id placeholder is filled in by Stripe before redirect.
            success_url=(
                f"{settings.public_app_url}/find-lawyer/{request_id}/paid"
                "?session_id={CHECKOUT_SESSION_ID}"
            ),
            cancel_url=f"{settings.public_app_url}/find-lawyer/{request_id}",
            client_reference_id=request_id,
            metadata={"professional_search_id": request_id},
            # Capture an email so we can pre-fill signup post-purchase.
            customer_email=None,  # let Stripe collect — pre-filling requires we have it
        )
    except Exception as exc:
        logger.exception("stripe checkout creation failed for %s", request_id)
        raise HTTPException(status_code=502, detail=f"Stripe error: {exc}")

    # Optimistic write — webhook is the source of truth, but recording the
    # session id here lets us recover if the webhook ever lags.
    row.stripe_session_id = checkout.id

    # Stage 2 dispatch — fire enrichment as a background task IFF we have
    # firms_data to enrich AND we haven't already kicked off enrichment.
    # This runs in parallel with the user's Stripe session; by the time
    # they finish entering card details (~30-60s), enrichment is usually
    # 1/3 done. The /paid page polls /enrichment-status to update the UI
    # when it lands. Idempotent: re-clicks don't re-dispatch.
    if (
        row.firms_data
        and (row.enrichment_status or "idle") in ("idle", "failed")
    ):
        from compliance_os.web.services.enrichment_runner import run_enrichment_sync

        row.enrichment_status = "enriching"
        row.enrichment_started_at = _dt.datetime.utcnow()
        row.enrichment_error = None
        background_tasks.add_task(run_enrichment_sync, row.id)
        logger.info(
            "checkout: dispatched enrichment for %s (firms=%d)",
            request_id, len(row.firms_data),
        )

    session.commit()
    return {"url": checkout.url, "session_id": checkout.id}


@router.post("/{request_id}/sync-stripe-session", response_model=SearchResponse)
def sync_stripe_session(
    request_id: str,
    session_id: str,
    db: Session = Depends(get_session),
):
    """Webhook fallback — verify a Stripe Checkout Session directly and
    apply the same updates the webhook would have, idempotently.

    Triggered from the /paid page's mount-effect when the URL has
    session_id in it (every Stripe success_url redirect does). Stripe's
    webhook delivery is occasionally flaky — when the Stripe dashboard
    webhook config drifts (wrong domain, expired secret, disabled
    endpoint), payments succeed but `paid_at` never gets set, leaving
    the user staring at "Confirming your payment…" forever even though
    the money was already taken.

    This endpoint short-circuits that. It calls Stripe directly with our
    server-side secret key, confirms the session was paid for THIS search
    (`client_reference_id == request_id`), then commits paid_at + the
    saved customer email exactly like the webhook handler. Idempotent —
    if paid_at is already set we just return the current row state.
    """
    row = db.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        raise HTTPException(status_code=404, detail="Search request not found")
    if row.paid_at is not None:
        # Already paid (probably by a webhook that did fire, or a prior
        # call to this endpoint). No-op.
        return _serialize(row)

    stripe = _stripe()
    try:
        sess = stripe.checkout.Session.retrieve(session_id)
    except Exception as exc:
        logger.exception("sync-stripe-session: retrieve failed for %s", request_id)
        raise HTTPException(status_code=502, detail=f"Stripe error: {exc}")

    # Two safety checks that the webhook handler also enforces:
    sess_request_id = (
        sess.get("client_reference_id")
        if isinstance(sess, dict)
        else getattr(sess, "client_reference_id", None)
    )
    if sess_request_id != request_id:
        logger.warning(
            "sync-stripe-session: client_reference_id mismatch — "
            "session %s belongs to %s, not %s",
            session_id, sess_request_id, request_id,
        )
        raise HTTPException(
            status_code=400,
            detail="Stripe session does not belong to this search request",
        )
    payment_status = (
        sess.get("payment_status")
        if isinstance(sess, dict)
        else getattr(sess, "payment_status", None)
    )
    if payment_status != "paid":
        raise HTTPException(
            status_code=409,
            detail=f"Stripe session is not paid (payment_status={payment_status})",
        )

    # All checks pass — apply the webhook's writes.
    row.paid_at = _dt.datetime.utcnow()
    row.stripe_session_id = (
        sess.get("id") if isinstance(sess, dict) else getattr(sess, "id", session_id)
    )
    customer_details = (
        sess.get("customer_details") if isinstance(sess, dict)
        else getattr(sess, "customer_details", None)
    )
    if customer_details:
        email = (
            customer_details.get("email")
            if isinstance(customer_details, dict)
            else getattr(customer_details, "email", None)
        )
        if email and not row.stripe_customer_email:
            row.stripe_customer_email = email
    db.commit()
    db.refresh(row)
    logger.info("sync-stripe-session: marked %s paid via API fallback", request_id)
    return _serialize(row)


@router.get(
    "/{request_id}/enrichment-status",
    response_model=EnrichmentStatusResponse,
)
def get_enrichment_status(
    request_id: str,
    session: Session = Depends(get_session),
):
    """Lightweight polling endpoint for the /paid page.

    Returns just the enrichment lifecycle state + a per-firm completion
    counter so the UI can render a "X of Y firms enriched" progress hint
    during the ~2-3min enrichment window. Avoids dragging the full
    firms_data + tier_report payload across the wire on every poll.
    """
    row = session.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        raise HTTPException(status_code=404, detail="Search request not found")

    firms_total = len(row.firms_data or [])
    firms_enriched = sum(
        1
        for f in (row.firms_data or [])
        if isinstance(f, dict) and f.get("_enriched_at")
    )
    return EnrichmentStatusResponse(
        status=row.enrichment_status or "idle",
        started_at=(
            row.enrichment_started_at.isoformat()
            if row.enrichment_started_at else None
        ),
        completed_at=(
            row.enrichment_completed_at.isoformat()
            if row.enrichment_completed_at else None
        ),
        firms_enriched=firms_enriched,
        firms_total=firms_total,
    )


@router.post("/stripe-webhook")
async def stripe_webhook(
    request: Request,
    stripe_signature: str | None = Header(None, alias="stripe-signature"),
    db: Session = Depends(get_session),
):
    """Handle `checkout.session.completed` — the canonical 'paid' signal.

    Verified via `STRIPE_WEBHOOK_SECRET`. Idempotent: re-deliveries that
    re-process the same session are a no-op.
    """
    payload = await request.body()
    if not settings.stripe_webhook_secret:
        raise HTTPException(status_code=503, detail="Stripe webhook secret not configured")
    if not stripe_signature:
        raise HTTPException(status_code=400, detail="Missing stripe-signature header")

    stripe = _stripe()
    try:
        stripe.Webhook.construct_event(
            payload, stripe_signature, settings.stripe_webhook_secret
        )
    except Exception as exc:
        logger.warning("stripe webhook signature verification failed: %s", exc)
        raise HTTPException(status_code=400, detail=f"Invalid signature: {exc}")

    # stripe-python 15.x's StripeObject no longer inherits from dict and
    # `to_dict_recursive` was removed, so the previous shim left `obj`
    # as a `Session` whose `.get(...)` raises AttributeError. We already
    # have the verified raw bytes — parse them as JSON for a plain dict.
    event = json.loads(payload)
    event_type = event["type"]
    obj = event["data"]["object"]

    # Subscription lifecycle events — drive the SubscriptionRow mirror.
    # Note: customer.subscription.created arrives near-simultaneously with
    # checkout.session.completed (mode=subscription); both can create the
    # row, so the upsert helper is idempotent.
    if event_type in (
        "customer.subscription.created",
        "customer.subscription.updated",
        "customer.subscription.deleted",
    ):
        _handle_subscription_event(event_type, obj, db)
        return {"received": True, "kind": "subscription_event", "type": event_type}

    if event_type != "checkout.session.completed":
        return {"received": True, "ignored": event_type}

    # Branch on Checkout mode: subscription signup vs. one-time payment.
    if obj.get("mode") == "subscription":
        _handle_subscription_checkout_completed(obj, db, stripe)
        return {"received": True, "kind": "subscription_checkout_completed"}

    # One-time payment → existing lawyer-search flow. The variable is
    # named `sess` from here for git-blame stability with the historical
    # implementation below.
    sess = obj

    request_id = sess.get("client_reference_id") or (sess.get("metadata") or {}).get(
        "professional_search_id"
    )
    if not request_id:
        # Operational red flag — every checkout we create sets this. If
        # we receive a verified webhook without it, either someone is
        # creating Sessions outside our flow against the same account,
        # or our own create_checkout_session call dropped client_reference_id.
        logger.error(
            "STRIPE_WEBHOOK_ALERT: verified checkout.session.completed with no "
            "client_reference_id and no metadata.professional_search_id. "
            "session_id=%s amount_total=%s customer_email=%s",
            sess.get("id"),
            sess.get("amount_total"),
            (sess.get("customer_details") or {}).get("email"),
        )
        return {"received": True, "ignored": "no client_reference_id"}

    row = db.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        # The row was deleted between checkout and webhook delivery (rare —
        # we don't currently delete rows). Log structured so we can grep
        # alerts; refunds may be needed if the customer was actually charged.
        logger.error(
            "STRIPE_WEBHOOK_ALERT: webhook for non-existent search row. "
            "request_id=%s session_id=%s amount_total=%s customer_email=%s",
            request_id,
            sess.get("id"),
            sess.get("amount_total"),
            (sess.get("customer_details") or {}).get("email"),
        )
        return {"received": True, "ignored": "unknown row"}

    if row.paid_at is not None:
        # Idempotency — Stripe retries successful webhooks if our 2xx is missed.
        return {"received": True, "already_paid": True}

    row.paid_at = _dt.datetime.utcnow()
    row.stripe_session_id = sess.get("id")
    customer_email = (sess.get("customer_details") or {}).get("email") or sess.get(
        "customer_email"
    )
    if customer_email:
        row.stripe_customer_email = customer_email

        # If a user with that email already exists, link the search now —
        # saves a step in the post-purchase signup flow.
        existing = (
            db.query(UserRow).filter(UserRow.email == customer_email).first()
        )
        if existing is not None and row.user_id is None:
            row.user_id = existing.id

    # If the commit fails AFTER Stripe has already charged the customer,
    # the user is in the worst state: they paid but our DB never reflects
    # it. Stripe will retry the webhook, but if the failure is persistent
    # (constraint violation, schema drift), every retry will fail. Log
    # an explicit ALERT so we catch it in observability and can manually
    # reconcile (mark paid_at by hand, or refund). The HTTP 500 we then
    # return causes Stripe to retry — preserving the canonical recovery
    # path for transient failures.
    try:
        db.commit()
    except Exception as exc:
        db.rollback()
        logger.error(
            "STRIPE_WEBHOOK_ALERT: db.commit failed AFTER customer was charged. "
            "request_id=%s session_id=%s amount_total=%s customer_email=%s exc=%s",
            request_id,
            sess.get("id"),
            sess.get("amount_total"),
            customer_email,
            exc,
        )
        # Re-raise so Stripe sees a 5xx and retries. If the failure is
        # transient the next retry succeeds; if persistent we'll have
        # the alert log + the original Stripe charge to manually resolve.
        raise HTTPException(status_code=500, detail="commit failed; will retry")
    logger.info("stripe webhook: marked search %s paid", request_id)

    # NOTE: trials are no longer silently auto-created on $15 payment.
    # With saved-card now in place (`setup_future_usage`), an auto-attach
    # would mean a surprise $20/mo charge after 30 days — exactly the
    # negative-option pattern we want to avoid. Trial creation now flows
    # through the explicit "Start Pro free for 30 days" CTA on the paid
    # page → POST /{request_id}/start-pro-trial.

    return {"received": True, "paid": True}


# ---------------------------- Subscription helpers ----------------------------


def _stripe_obj_to_dict(obj: Any) -> dict:
    """Coerce a stripe-python `StripeObject` (15.x) to a plain dict.

    `StripeObject.to_dict()` returns a top-level dict but nested values
    can still be `StripeObject`s — and StripeObject no longer inherits
    from dict in 15.x, so callers that do `obj.get(...)` raise
    AttributeError. JSON round-trip flattens everything; `default=str`
    covers any non-JSON-native values Stripe might surface.
    """
    if isinstance(obj, dict):
        return obj
    base = obj.to_dict() if hasattr(obj, "to_dict") else obj
    return json.loads(json.dumps(base, default=str))


def _to_dt_or_none(epoch: int | None) -> _dt.datetime | None:
    if not epoch:
        return None
    return _dt.datetime.utcfromtimestamp(int(epoch))


def _upsert_subscription_from_stripe(
    sub_obj: dict,
    user_id: str | None,
    db: Session,
) -> "SubscriptionRow | None":
    """Mirror a Stripe Subscription object into our subscriptions table.

    Idempotent — keyed on stripe_subscription_id. Returns the row, or
    None if we cannot resolve a user_id (orphan subscription — log and
    skip rather than create rows we can't attribute).
    """
    from compliance_os.web.models.auth import SubscriptionRow as _SubRow
    from compliance_os.web.services.subscription_service import derive_tier

    sub_id = sub_obj.get("id")
    if not sub_id:
        logger.error("STRIPE_WEBHOOK_ALERT: subscription event with no id")
        return None

    row = (
        db.query(_SubRow)
        .filter(_SubRow.stripe_subscription_id == sub_id)
        .first()
    )

    # Resolve user_id — prefer caller-supplied, then metadata.user_id on
    # the sub itself. If neither, the row is unattributable.
    if user_id is None:
        meta = sub_obj.get("metadata") or {}
        user_id = meta.get("user_id")
    if user_id is None and row is not None:
        user_id = row.user_id  # already known from earlier event
    if user_id is None:
        logger.error(
            "STRIPE_WEBHOOK_ALERT: subscription %s has no resolvable user_id "
            "(no metadata.user_id, no existing row). Skipping.",
            sub_id,
        )
        return None

    status = sub_obj.get("status") or "incomplete"
    is_canceled = status in ("canceled", "incomplete_expired", "unpaid")
    new_values = {
        "user_id": user_id,
        "stripe_customer_id": sub_obj.get("customer"),
        "stripe_subscription_id": sub_id,
        "stripe_price_id": _extract_price_id(sub_obj),
        "status": status,
        "tier": derive_tier(status),
        "current_period_start": _to_dt_or_none(sub_obj.get("current_period_start")),
        "current_period_end": _to_dt_or_none(sub_obj.get("current_period_end")),
        "trial_end": _to_dt_or_none(sub_obj.get("trial_end")),
        "cancel_at_period_end": bool(sub_obj.get("cancel_at_period_end")),
        "canceled_at": (
            _to_dt_or_none(sub_obj.get("canceled_at"))
            if is_canceled
            else None
        ),
    }

    if row is None:
        row = _SubRow(**new_values)
        db.add(row)
    else:
        for k, v in new_values.items():
            setattr(row, k, v)
    return row


def _extract_price_id(sub_obj: dict) -> str | None:
    """Pull the first price id off a Stripe Subscription object.

    Subscriptions support multiple line items in theory, but our Pro
    product is a single price. Defensive: return None if shape changes.
    """
    items = (sub_obj.get("items") or {}).get("data") or []
    if not items:
        return None
    price = items[0].get("price") or {}
    return price.get("id")


def _handle_subscription_event(event_type: str, sub_obj: dict, db: Session) -> None:
    """customer.subscription.{created,updated,deleted} → upsert mirror row."""
    row = _upsert_subscription_from_stripe(sub_obj, user_id=None, db=db)
    try:
        db.commit()
    except Exception:
        db.rollback()
        logger.exception(
            "STRIPE_WEBHOOK_ALERT: subscription %s commit failed for %s",
            event_type,
            sub_obj.get("id"),
        )
        raise HTTPException(status_code=500, detail="commit failed; will retry")
    if row is not None:
        logger.info(
            "stripe webhook: subscription %s %s tier=%s status=%s",
            row.stripe_subscription_id,
            event_type,
            row.tier,
            row.status,
        )


def _handle_subscription_checkout_completed(sess: dict, db: Session, stripe) -> None:
    """checkout.session.completed for mode=subscription.

    Stripe also fires customer.subscription.created at the same time, so
    most of the time the SubscriptionRow is already in place by the time
    this handler runs. But ordering is not guaranteed, so we fetch the
    Subscription and upsert defensively.
    """
    sub_id = sess.get("subscription")
    if not sub_id:
        logger.error(
            "STRIPE_WEBHOOK_ALERT: subscription checkout %s has no subscription id",
            sess.get("id"),
        )
        return

    user_id = sess.get("client_reference_id") or (sess.get("metadata") or {}).get("user_id")
    try:
        full = stripe.Subscription.retrieve(sub_id)
        sub_obj = _stripe_obj_to_dict(full)
    except Exception:
        logger.exception(
            "STRIPE_WEBHOOK_ALERT: failed to fetch subscription %s", sub_id
        )
        return

    _upsert_subscription_from_stripe(sub_obj, user_id, db)
    try:
        db.commit()
    except Exception:
        db.rollback()
        logger.exception(
            "STRIPE_WEBHOOK_ALERT: subscription_checkout commit failed for %s", sub_id
        )
        raise HTTPException(status_code=500, detail="commit failed; will retry")


def _maybe_attach_trial(
    user: UserRow,
    db: Session,
    *,
    stripe,
    customer_id: str | None,
    source: str,
) -> None:
    """Create a 30-day Pro trial for `user` if they don't already have one.

    No-op if:
      * user already has any active subscription (paid or trialing)
      * STRIPE_PRO_PRICE_ID is not configured (graceful skip — the trial
        is a benefit, not a payment, so missing config shouldn't crash
        the calling search-payment flow)

    Stripe will fire customer.subscription.created back to our webhook,
    which is what actually persists the SubscriptionRow. This function
    just initiates the Stripe-side state.
    """
    from compliance_os.web.services.subscription_service import get_active_subscription

    if not settings.stripe_pro_price_id:
        return  # Pro not configured — skip trial creation.
    if get_active_subscription(user, db) is not None:
        return  # Already entitled.

    # Resolve a Stripe Customer for this user — reuse the one from the
    # one-time charge if present, otherwise create one keyed by email.
    if not customer_id:
        try:
            c = stripe.Customer.create(
                email=user.email,
                metadata={"user_id": user.id},
            )
            customer_id = c.id if hasattr(c, "id") else c.get("id")
        except Exception:
            logger.exception(
                "trial attach: failed to create Stripe customer for %s", user.id
            )
            return

    trial_end = int(
        (_dt.datetime.utcnow() + _dt.timedelta(days=30)).timestamp()
    )
    try:
        stripe.Subscription.create(
            customer=customer_id,
            items=[{"price": settings.stripe_pro_price_id}],
            trial_end=trial_end,
            # If the user never adds a card during the trial, Stripe will
            # cancel the sub at trial_end rather than fail to charge.
            trial_settings={"end_behavior": {"missing_payment_method": "cancel"}},
            payment_settings={"save_default_payment_method": "on_subscription"},
            metadata={
                "user_id": user.id,
                "trial_source": source,
            },
        )
    except Exception:
        logger.exception(
            "trial attach: Stripe.Subscription.create failed for user %s", user.id
        )
        return
    logger.info("trial attach: 30-day Pro trial created for user %s (source=%s)",
                user.id, source)


@router.post("/{request_id}/start-pro-trial")
def start_pro_trial(
    request_id: str,
    authorization: str | None = Header(None),
    db: Session = Depends(get_session),
):
    """Opt-in: start a 30-day Pro trial that auto-renews at \$20/mo.

    Called from the paid page's "Keep Pro free for 30 days" CTA. The
    user has just paid \$15 for a one-off search — that checkout saved
    their card via `setup_future_usage`. We use that saved card to
    create a Stripe subscription with a 30-day trial; Stripe auto-bills
    the standard Pro price (\$20/mo) at trial end unless the user
    cancels via the billing portal.

    Idempotent: if the user already has any active subscription
    (paid or trialing), returns ok without creating a duplicate.
    """
    payload = get_bearer_payload(authorization, db)
    user = db.query(UserRow).filter(UserRow.id == payload["user_id"]).first()
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")

    row = db.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        raise HTTPException(status_code=404, detail="Search request not found")
    if row.user_id != user.id:
        raise HTTPException(
            status_code=403,
            detail="This search isn't claimed under your account.",
        )
    if row.paid_at is None or not row.stripe_session_id:
        raise HTTPException(
            status_code=409,
            detail="Need a paid one-off search before starting a Pro trial.",
        )

    from compliance_os.web.services.subscription_service import get_active_subscription
    if get_active_subscription(user, db) is not None:
        return {"ok": True, "already_active": True}

    if not settings.stripe_pro_price_id:
        raise HTTPException(
            status_code=503,
            detail="Pro plan not configured on server (STRIPE_PRO_PRICE_ID).",
        )

    stripe = _stripe()

    # Pull the saved Stripe Customer + payment method from the $15 payment.
    # Sessions older than the setup_future_usage change won't have a saved
    # card; in that case we surface a clear error so the UI can fall back
    # to the standard /api/subscription/checkout flow (re-enter card).
    try:
        sess = stripe.checkout.Session.retrieve(
            row.stripe_session_id,
            expand=["payment_intent.payment_method"],
        )
    except Exception as exc:
        logger.exception("stripe session retrieve failed for %s", request_id)
        raise HTTPException(status_code=502, detail=f"Stripe error: {exc}")

    customer_id = (
        sess.get("customer") if isinstance(sess, dict) else getattr(sess, "customer", None)
    )
    if not customer_id:
        raise HTTPException(
            status_code=409,
            detail=(
                "No saved card on this purchase — start a Pro subscription via "
                "/pricing instead (you'll re-enter card details)."
            ),
        )

    _maybe_attach_trial(
        user,
        db,
        stripe=stripe,
        customer_id=customer_id,
        source="post_search_cta",
    )

    return {
        "ok": True,
        "trial_days": 30,
        "renewal_price_cents": 2000,  # informational — actual amount comes from Stripe price
    }


@router.post("/{request_id}/claim", response_model=SearchResponse)
def claim_search(
    request_id: str,
    authorization: str | None = Header(None),
    db: Session = Depends(get_session),
):
    """Link a paid search row to the authenticated user.

    Used after pay-first checkout: user signs up (or logs in), then their
    client calls this with their bearer token. The row's `user_id` is set,
    making it appear under "My searches" in the dashboard.
    """
    payload = get_bearer_payload(authorization, db)
    user = db.query(UserRow).filter(UserRow.id == payload["user_id"]).first()
    if user is None:
        raise HTTPException(status_code=401, detail="User not found")

    row = db.get(ProfessionalSearchRequestRow, request_id)
    if row is None:
        raise HTTPException(status_code=404, detail="Search request not found")
    if row.paid_at is None:
        raise HTTPException(
            status_code=402,
            detail="Cannot claim an unpaid search — complete checkout first.",
        )

    # If already claimed by this same user, treat as idempotent success.
    # If claimed by SOMEONE ELSE, reject — even with a valid token.
    if row.user_id is not None and row.user_id != user.id:
        raise HTTPException(
            status_code=409,
            detail="This search has already been claimed by another account.",
        )

    # Defense-in-depth: require the claimer's email to match the email
    # Stripe captured at checkout. UUIDs are unguessable but the search
    # status URL is shareable, and a Referer leak or screenshot would
    # otherwise let any authed account claim a stranger's paid search.
    # Only enforce when Stripe gave us an email (older rows may be null).
    if (
        row.stripe_customer_email
        and row.stripe_customer_email.strip().lower() != (user.email or "").strip().lower()
    ):
        raise HTTPException(
            status_code=403,
            detail=(
                "Email mismatch: this search was paid for by a different "
                "email. Sign in (or sign up) with the email you used at "
                "checkout to claim it."
            ),
        )

    row.user_id = user.id

    # Skip-onboarding shortcut: if the search came with a substantive
    # brief or supporting docs and isn't already case-attached, mint a
    # case for it now. The user will be redirected straight to the case
    # page (post-claim flow) instead of being forced through the empty
    # discovery wizard. They can always fill the wizard later for richer
    # follow-ups — it remains available, just not blocking.
    #
    # Threshold rationale: 400 chars is "above the 200-char minimum + a
    # bit more than the floor", which signals real intent. Any uploaded
    # files are an even stronger signal — the user already gathered
    # documents, so they clearly know their situation.
    BRIEF_RICHNESS_CHARS = 400
    has_uploads = bool((row.uploaded_notes or "").strip())
    brief_is_rich = (
        len((row.case_brief or "").strip()) >= BRIEF_RICHNESS_CHARS
        or has_uploads
    )
    if row.case_id is None and brief_is_rich:
        from compliance_os.web.models.tables import CaseRow as _CaseRow

        # Map vertical → workflow_type so the case shows the right track
        # in dashboards and CTAs. Verticals without a clear track default
        # to immigration (the most common path through find-lawyer).
        _VERTICAL_TO_WORKFLOW = {
            "immigration_attorney": "immigration",
            "immigration_eb5": "immigration",
            "tax_attorney": "tax",
            "cpa": "tax",
            "caa": "tax",
            "corporate_attorney": "corporate",
            "bank": "",
        }
        case = _CaseRow(
            user_id=user.id,
            workflow_type=_VERTICAL_TO_WORKFLOW.get(row.vertical, "immigration"),
            # `status="discovery"` is the default. Leave it — if the user
            # opens the discovery wizard later it picks up correctly.
        )
        db.add(case)
        db.flush()  # need case.id before assigning back
        row.case_id = case.id
        logger.info(
            "claim auto-created case %s for search %s (brief=%dch uploads=%s)",
            case.id, request_id, len((row.case_brief or "").strip()), has_uploads,
        )

    db.commit()
    db.refresh(row)

    # Trial creation is no longer auto-attached on claim. Users opt in
    # via the explicit "Start Pro free for 30 days" CTA on the paid page,
    # which calls POST /{request_id}/start-pro-trial. Keeps the claim
    # action transactionally simple (just links search → user) and
    # avoids any surprise auto-billing once the saved card lands.

    # Caller is now the authenticated owner — return full PII.
    return _serialize(row, include_pii=True)


# NOTE: `GET /mine/list` is defined above near the top of the router
# (next to `GET /{request_id}`) so the static-prefix route appears
# before the catch-all in declaration order. The duplicate that used
# to live here was removed during the code-review fixes.
